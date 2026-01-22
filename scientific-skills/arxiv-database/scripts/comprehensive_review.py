"""
Comprehensive Paper Review Workflow (CPRW)

⚠️  DEPRECATED - This script is deprecated and will be removed in a future version.

**MIGRATION PATH:**
- For structure extraction → Use paper_structure_extractor.py
- For peer review → Use the 'peer-review' skill
- For validation → Use the 'paper-validator' skill
- For literature comparison → Use the 'literature-review' skill

This script remains available for backward compatibility but is no longer maintained.

⚠️  EXPERIMENTAL FEATURE - This is a template tool with significant limitations.

A 7-phase systematic framework template for peer review of scientific papers:
1. Paper Input & Parsing - Support for DOCX, PDF, arXiv ID
2. Literature Search & Collection - Multi-source retrieval
3. Literature Analysis & Comparison - Structured analysis
4. Comprehensive Evaluation - Scoring and assessment
5. Peer Review - 7-stage review framework
6. Literature Comparison Report - Gap analysis
7. Final Review Report - Comprehensive output

IMPORTANT LIMITATIONS:
- This tool generates a review TEMPLATE, not actual peer review
- Evaluations and scores are structural placeholders
- Actual content analysis requires human expert review
- Use this for parsing and structure, not for automated evaluation

Dependencies:
    pip install pypdf>=3.0.0 python-docx>=1.0.0
    Or: pip install -e .

Usage:
    python comprehensive_review.py --input /path/to/paper.docx --output review.json
    python comprehensive_review.py --arxiv-id 2401.12345 --output review.json
    python comprehensive_review.py --input /path/to/paper.pdf --output review.json --verbose
"""

import argparse
import warnings

# Issue deprecation warning
warnings.warn(
    "\n\n"
    "=" * 70 + "\n"
    "DEPRECATION WARNING: comprehensive_review.py is deprecated.\n"
    "\n"
    "This script will be removed in a future version.\n"
    "\n"
    "Migration path:\n"
    "  - Structure extraction → Use paper_structure_extractor.py\n"
    "  - Peer review → Use 'peer-review' skill\n"
    "  - Validation → Use 'paper-validator' skill\n"
    "  - Literature comparison → Use 'literature-review' skill\n"
    "\n"
    "See SKILL.md for details on the new workflow.\n"
    "=" * 70 + "\n",
    DeprecationWarning,
    stacklevel=2
)
import json
import logging
import re
import sys
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional
import urllib.request
import urllib.error
import urllib.parse

try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning(
        "pypdf not available. PDF parsing will be disabled. "
        "Install with: pip install pypdf>=3.0.0"
    )

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning(
        "python-docx not available. DOCX parsing will be disabled. "
        "Install with: pip install python-docx>=1.0.0"
    )


logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class PaperSection:
    """Represents a section of a paper."""
    title: str
    content: str
    level: int = 1


@dataclass
class PaperReference:
    """Represents a reference from a paper."""
    id: str
    text: str
    arxiv_id: Optional[str] = None
    doi: Optional[str] = None


@dataclass
class PaperFigure:
    """Represents a figure in a paper."""
    number: str
    caption: str
    page_estimate: Optional[int] = None


@dataclass
class PaperTable:
    """Represents a table in a paper."""
    number: str
    caption: str


@dataclass
class PaperMetadata:
    """Complete metadata for a paper."""
    title: str
    authors: list[str] = field(default_factory=list)
    abstract: str = ""
    categories: list[str] = field(default_factory=list)
    arxiv_id: Optional[str] = None
    doi: Optional[str] = None
    pdf_url: Optional[str] = None
    published: Optional[str] = None
    updated: Optional[str] = None
    comment: Optional[str] = None
    journal_ref: Optional[str] = None
    source_type: str = "unknown"


@dataclass
class StructuredPaperData:
    """Complete structured paper data."""
    metadata: PaperMetadata
    sections: dict[str, PaperSection] = field(default_factory=dict)
    figures: list[PaperFigure] = field(default_factory=list)
    tables: list[PaperTable] = field(default_factory=list)
    references: list[PaperReference] = field(default_factory=list)


class PaperParser(ABC):
    """Abstract base class for paper parsers."""

    @abstractmethod
    def parse(self, source: str) -> StructuredPaperData:
        """Parse a paper from the given source."""
        pass


class DocxParser:
    """Parser for DOCX documents."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose

    def parse(self, file_path: str) -> StructuredPaperData:
        """Parse a DOCX document."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx>=1.0.0 "
                "or install all dependencies: pip install -e ."
            )

        doc = docx.Document(file_path)

        metadata = PaperMetadata(
            title=self._extract_title(doc),
            source_type="docx"
        )

        full_text = []
        sections = {}

        current_section = "Introduction"
        section_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text.append(text)

            section_header = self._detect_section_header(text)
            if section_header:
                if section_content:
                    sections[current_section] = PaperSection(
                        title=current_section,
                        content="\n".join(section_content)
                    )
                current_section = section_header
                section_content = []
            else:
                section_content.append(text)

        if section_content:
            sections[current_section] = PaperSection(
                title=current_section,
                content="\n".join(section_content)
            )

        abstract = sections.get("Abstract", PaperSection(title="Abstract", content=""))
        metadata.abstract = abstract.content if abstract else ""

        return StructuredPaperData(
            metadata=metadata,
            sections=sections,
            figures=self._extract_figures(doc),
            tables=self._extract_tables(doc),
            references=self._extract_references(doc)
        )

    def _extract_title(self, doc) -> str:
        """Extract title from document."""
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text.strip()
            if len(first_para) < 300:
                return first_para
        return "Unknown Title"

    def _detect_section_header(self, text: str) -> Optional[str]:
        """Detect if text is a section header."""
        patterns = [
            r'^(abstract)$',
            r'^1\.?\s*(introduction|overview)',
            r'^2\.?\s*(related work|background)',
            r'^3\.?\s*(method|approach|methodology)',
            r'^4\.?\s*(experiment|evaluation)',
            r'^5\.?\s*(result|results)',
            r'^6\.?\s*(discussion|conclusion)',
            r'^(acknowledgments?)',
            r'^(references|bibliography)',
        ]

        text_lower = text.lower().strip()
        for pattern in patterns:
            if re.match(pattern, text_lower):
                return text.title()

        if len(text) < 100 and text.isupper():
            return text.title()

        return None

    def _extract_figures(self, doc) -> list[PaperFigure]:
        """Extract figures from document."""
        figures = []
        pattern = r'(?:Figure|Fig\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n|$)'

        for para in doc.paragraphs:
            text = para.text.strip()
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                figures.append(PaperFigure(
                    number=match.group(1),
                    caption=match.group(2)[:200]
                ))

        return figures

    def _extract_tables(self, doc) -> list[PaperTable]:
        """Extract tables from document."""
        tables = []
        pattern = r'(?:Table|Tab\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n|$)'

        for para in doc.paragraphs:
            text = para.text.strip()
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                tables.append(PaperTable(
                    number=match.group(1),
                    caption=match.group(2)[:200]
                ))

        return tables

    def _extract_references(self, doc) -> list[PaperReference]:
        """Extract references from document."""
        references = []
        ref_pattern = r'\[(\d+)\]\s*(.+?)(?=\n\s*\[|\n\s*\d+\.|$)'
        in_references = False

        for para in doc.paragraphs:
            text = para.text.strip()

            if re.match(r'^(references|bibliography)$', text.lower()):
                in_references = True
                continue

            if in_references:
                match = re.match(ref_pattern, text, re.DOTALL)
                if match:
                    ref_text = match.group(2).strip()
                    arxiv_id = self._extract_arxiv_id(ref_text)
                    references.append(PaperReference(
                        id=match.group(1),
                        text=ref_text[:500],
                        arxiv_id=arxiv_id
                    ))

        return references

    def _extract_arxiv_id(self, text: str) -> Optional[str]:
        """Extract arXiv ID from text."""
        match = re.search(r'arXiv[:\.]?\s*(\d+\.\d+)', text, re.IGNORECASE)
        return match.group(1) if match else None


class PdfParser:
    """Parser for PDF documents using pypdf."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose

    def parse(self, file_path: str) -> StructuredPaperData:
        """Parse a PDF document."""
        if not PDF_AVAILABLE:
            raise ImportError(
                "pypdf is required for PDF parsing. "
                "Install with: pip install pypdf>=3.0.0 "
                "or install all dependencies: pip install -e ."
            )

        reader = pypdf.PdfReader(file_path)
        full_text = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)

        text = "\n\n".join(full_text)

        sections = self._extract_sections(text)
        metadata = self._extract_metadata(text, file_path)
        metadata.source_type = "pdf"

        abstract = sections.get("Abstract", PaperSection(title="Abstract", content=""))
        metadata.abstract = abstract.content if abstract else ""

        return StructuredPaperData(
            metadata=metadata,
            sections=sections,
            figures=self._extract_figures(text),
            tables=self._extract_tables(text),
            references=self._extract_references(text)
        )

    def _extract_sections(self, text: str) -> dict[str, PaperSection]:
        """Extract sections from text."""
        sections = {}

        section_patterns = {
            "abstract": r'(?i)(?:^|\n)(abstract)\s*[:\n](.+?)(?=\n\s*(?:1\.?|introduction|background|related work))',
            "introduction": r'(?i)(?:^|\n)(1\s*\.?\s*(?:introduction)?)\s*[:\n](.+?)(?=\n\s*(?:2\.?|related work|background))',
            "related_work": r'(?i)(?:^|\n)(2\s*\.?\s*(?:related work|background)?)\s*[:\n](.+?)(?=\n\s*(?:3\.?|method|approach|methodology))',
            "methodology": r'(?i)(?:^|\n)(3\s*\.?\s*(?:method|approach|methodology)?)\s*[:\n](.+?)(?=\n\s*(?:4\.?|experiment|evaluation|results|discussion))',
            "experiments": r'(?i)(?:^|\n)(4\s*\.?\s*(?:experiment|evaluation|experiments)?)\s*[:\n](.+?)(?=\n\s*(?:5\.?|result|discussion|conclusion))',
            "results": r'(?i)(?:^|\n)(5\s*\.?\s*(?:result|results|discussion)?)\s*[:\n](.+?)(?=\n\s*(?:6\.?|conclusion|future|limitation))',
            "conclusion": r'(?i)(?:^|\n)(6\s*\.?\s*(?:conclusion|conclusions|discussion)?)\s*[:\n](.+?)(?=\n\s*(?:references|bibliography|$))',
            "acknowledgments": r'(?i)(?:^|\n)(acknowledgments?)\s*[:\n](.+?)(?=\n\s*(?:references|bibliography|$))',
        }

        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.DOTALL)
            if match:
                content = match.group(2).strip()
                if len(content) > 10000:
                    content = content[:10000] + "..."
                sections[section_name] = PaperSection(
                    title=section_name.title(),
                    content=content
                )

        return sections

    def _extract_metadata(self, text: str, file_path: str) -> PaperMetadata:
        """Extract metadata from PDF text."""
        return PaperMetadata(
            title=self._extract_title(text),
            abstract="",
            source_type="pdf"
        )

    def _extract_title(self, text: str) -> str:
        """Extract title from text."""
        lines = text.split("\n")
        for line in lines[:10]:
            line = line.strip()
            if line and len(line) > 10 and len(line) < 300:
                if not line.startswith("http"):
                    return line
        return Path(file_path).stem

    def _extract_figures(self, text: str) -> list[PaperFigure]:
        """Extract figures from text."""
        figures = []
        pattern = r'(?:Figure|Fig\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n\s*(?:Figure|Fig\.?|Table|$))'

        for i, match in enumerate(re.finditer(pattern, text, re.IGNORECASE)):
            figures.append(PaperFigure(
                number=match.group(1),
                caption=match.group(2).strip()[:200]
            ))

        return figures

    def _extract_tables(self, text: str) -> list[PaperTable]:
        """Extract tables from text."""
        tables = []
        pattern = r'(?:Table|Tab\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n\s*(?:Table|Tab\.?|Figure|$))'

        for i, match in enumerate(re.finditer(pattern, text, re.IGNORECASE)):
            tables.append(PaperTable(
                number=match.group(1),
                caption=match.group(2).strip()[:200]
            ))

        return tables

    def _extract_references(self, text: str) -> list[PaperReference]:
        """Extract references from text."""
        references = []
        ref_match = re.search(r'(?i)(?:references|bibliography)\s*[:\n](.+)', text, re.DOTALL)
        if not ref_match:
            return references

        ref_text = ref_match.group(1)
        bracket_pattern = r'\[(\d+)\]\s*(.+?)(?=\n\s*\[|\n\s*\d+\.|\Z)'

        for match in re.finditer(bracket_pattern, ref_text, re.DOTALL):
            ref_content = match.group(2).strip()
            arxiv_id = self._extract_arxiv_id(ref_content)
            references.append(PaperReference(
                id=match.group(1),
                text=ref_content[:500],
                arxiv_id=arxiv_id
            ))

        return references

    def _extract_arxiv_id(self, text: str) -> Optional[str]:
        """Extract arXiv ID from text."""
        match = re.search(r'arXiv[:\.]?\s*(\d+\.\d+)', text, re.IGNORECASE)
        return match.group(1) if match else None


class ArxivApiError(Exception):
    """Exception for arXiv API errors."""
    pass


class SemanticKeywordExtractor:
    """
    Intelligent keyword extraction from paper text.
    
    Uses NLP techniques to extract meaningful terms without hardcoding.
    """

    STOP_WORDS = {
        "a", "an", "the", "is", "are", "was", "were", "of", "in", "on", "at",
        "for", "with", "and", "or", "to", "from", "by", "that", "this", "these",
        "those", "which", "what", "how", "when", "where", "why", "can", "could",
        "would", "should", "may", "might", "will", "shall", "must", "have", "has",
        "been", "being", "do", "does", "did", "doing", "done", "make", "made",
        "using", "used", "use", "based", "approach", "method", "technique",
        "paper", "work", "study", "research", "propose", "proposed", "present",
        "presented", "show", "shown", "shows", "demonstrate", "demonstrated"
    }

    TECHNICAL_TERMS = {
        "software_engineering": [
            "software engineering", "code generation", "code synthesis", "program generation",
            "software development", "software architecture", "software testing", "refactoring",
            "requirements engineering", "design patterns", "software maintenance"
        ],
        "systems": [
            "embedded system", "real-time system", "cyber-physical system", "distributed system",
            "operating system", "concurrent system", "parallel system", "safety-critical system",
            "system architecture", "system design", "system integration"
        ],
        "formal_methods": [
            "formal verification", "model checking", "theorem proving", "formal specification",
            "temporal logic", "safety verification", "formal methods", "formal analysis"
        ],
        "programming_languages": [
            "programming language", "compiler", "type system", "semantics", "parser",
            "interpreter", "domain-specific language", "program analysis", "static analysis"
        ],
        "ai_ml": [
            "machine learning", "deep learning", "neural network", "transformer",
            "large language model", "reinforcement learning", "natural language processing",
            "computer vision", "agent-based system", "multi-agent"
        ],
        "mbse": [
            "model-based systems engineering", "model transformation", "model-driven development",
            "sysml", "uml", "architecture description", "aadl", "osate", "model-to-code"
        ]
    }

    def __init__(self, verbose: bool = False):
        self.verbose = verbose

    def extract(self, text: str) -> dict:
        """
        Extract keywords and identify research domain from text.
        
        Returns:
            dict with 'keywords', 'bigrams', 'domains', 'technical_terms'
        """
        text_lower = text.lower()
        words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9]+\b', text_lower)

        filtered_words = [w for w in words if w not in self.STOP_WORDS and len(w) > 2]

        word_freq = {}
        for word in filtered_words:
            word_freq[word] = word_freq.get(word, 0) + 1

        keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:30]
        keywords = [w for w, _ in keywords]

        bigrams = self._extract_bigrams(text_lower)
        bigrams = sorted(set(bigrams), key=lambda x: x[1], reverse=True)[:15]
        bigrams = [b for b, _ in bigrams]

        domains = self._identify_domains(text_lower)

        technical_terms = []
        for domain, terms in self.TECHNICAL_TERMS.items():
            for term in terms:
                if term in text_lower:
                    technical_terms.append(term)

        technical_terms = list(set(technical_terms))

        return {
            "keywords": keywords,
            "bigrams": bigrams,
            "domains": domains,
            "technical_terms": technical_terms,
            "extracted_from": "title+abstract"
        }

    def _extract_bigrams(self, text: str) -> list[tuple[str, int]]:
        """Extract significant bigrams from text."""
        words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9]+\b', text)
        bigrams = []
        for i in range(len(words) - 1):
            if words[i] not in self.STOP_WORDS and words[i + 1] not in self.STOP_WORDS:
                bigram = f"{words[i]} {words[i + 1]}"
                bigrams.append((bigram, 1))
        bigram_counts = {}
        for bg, _ in bigrams:
            bigram_counts[bg] = bigram_counts.get(bg, 0) + 1
        return sorted(bigram_counts.items(), key=lambda x: x[1], reverse=True)

    def _identify_domains(self, text: str) -> list[str]:
        """Identify research domains based on content."""
        domains = []
        domain_scores = {}

        domain_patterns = {
            "software_engineering": [
                r"software\s+engineer", r"code\s+generation", r"software\s+development",
                r"software\s+architecture", r"software\s+testing"
            ],
            "embedded_systems": [
                r"embedded\s+system", r"real-time", r"cyber-physical",
                r"safety-critical", r"hardware"
            ],
            "formal_methods": [
                r"formal\s+verification", r"model\s+checking", r"theorem\s+proving",
                r"formal\s+specification", r"temporal\s+logic"
            ],
            "mbse": [
                r"model-based", r"model\s+driven", r"sysml", r"uml",
                r"architecture\s+description", r"aadl"
            ],
            "ai_ml": [
                r"machine\s+learning", r"deep\s+learning", r"neural\s+network",
                r"large\s+language\s+model", r"transformer", r"multi-agent",
                r"reinforcement\s+learn"
            ],
            "systems": [
                r"distributed\s+system", r"system\s+architecture", r"system\s+design"
            ]
        }

        for domain, patterns in domain_patterns.items():
            score = 0
            for pattern in patterns:
                if re.search(pattern, text):
                    score += 1
            if score > 0:
                domain_scores[domain] = score

        sorted_domains = sorted(domain_scores.items(), key=lambda x: x[1], reverse=True)
        domains = [d for d, _ in sorted_domains[:3]]

        return domains if domains else ["general"]

    def build_search_query(self, extracted_info: dict, original_query: str) -> str:
        """Build semantic search query from extracted information."""
        parts = [original_query]

        for term in extracted_info.get("technical_terms", []):
            if term.lower() not in original_query.lower():
                parts.append(term)

        for bigram, count in self._extract_bigrams(original_query.lower())[:5]:
            if bigram not in original_query.lower() and count > 1:
                parts.append(bigram)

        unique_parts = []
        seen = set()
        for p in parts:
            p_lower = p.lower()
            if p_lower not in seen:
                seen.add(p_lower)
                unique_parts.append(p)

        return " OR ".join(unique_parts)


class SemanticArxivSearcher:
    """
    Enhanced arXiv searcher with semantic search capabilities.
    
    Improvements over basic ArxivSearcher:
    1. Category-based filtering for relevant CS domains
    2. Multi-field search (title + abstract + all)
    3. Relevance scoring based on keyword matching
    4. Query expansion using research domain keywords
    """

    BASE_URL = "http://export.arxiv.org/api/query"

    ARXIV_CATEGORIES = {
        "software_engineering": ["cs.SE", "cs.PL"],
        "systems": ["cs.OS", "cs.DC", "cs.NI"],
        "programming_languages": ["cs.PL", "cs.SE"],
        "ai_ml": ["cs.AI", "cs.LG", "stat.ML"],
        "robotics": ["cs.RO", "cs.OS"],
        "embedded_systems": ["cs.OS", "cs.DC", "cs.SE"],
        "formal_methods": ["cs.LO", "cs.SE", "cs.FL"],
        "computer_architecture": ["cs.AR", "cs.DC"],
        "general_cs": ["cs.*"],
    }

    RELEVANT_CS_CATEGORIES = [
        "cs.SE", "cs.PL", "cs.OS", "cs.DC", "cs.NI", "cs.LO", "cs.FL",
        "cs.AI", "cs.LG", "cs.RO", "cs.AR", "cs.CV", "cs.CL", "cs.HC",
        "stat.ML", "math.OC"
    ]

    RELEVANT_KEYWORDS = {
        "software_engineering": [
            "software engineering", "code generation", "software development",
            "model-based", "MBSE", "software architecture", "requirement",
            "verification", "validation", "testing", "refactoring"
        ],
        "systems": [
            "embedded system", "real-time", "operating system", "distributed system",
            "concurrent", "parallel", "system architecture", "cyber-physical"
        ],
        "ai_ml": [
            "machine learning", "deep learning", "neural network", "transformer",
            "large language model", "LLM", "natural language processing"
        ],
        "formal_methods": [
            "formal verification", "model checking", "theorem proving",
            "formal specification", "temporal logic", "safety-critical"
        ],
        "programming_languages": [
            "programming language", "compiler", "type system", "semantics",
            "program synthesis", "program analysis", "static analysis"
        ],
    }

    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        self.request_delay = 3.0

    def search_by_id(self, arxiv_id: str) -> Optional[dict]:
        """Fetch paper by arXiv ID."""
        query = f"id_list={arxiv_id}"

        try:
            url = f"{self.BASE_URL}?{query}"
            response = self._make_request(url)

            if "entry" in response and response["entry"]:
                entry = response["entry"][0]
                entry["relevance_score"] = 1.0
                return entry

            return None

        except Exception as e:
            logger.error(f"Error fetching arXiv ID {arxiv_id}: {e}")
            return None

    def search(
        self,
        query: str,
        max_results: int = 20,
        categories: Optional[list[str]] = None,
        from_year: Optional[int] = None,
        sort_by: str = "relevance",
    ) -> list[dict]:
        """
        Semantic search for arXiv papers.
        
        Args:
            query: Search query string
            max_results: Maximum number of results
            categories: List of arXiv categories to filter (e.g., ["cs.SE", "cs.OS"])
            from_year: Filter papers from this year
            sort_by: Sort by "relevance", "submittedDate", or "lastUpdatedDate"
        """
        if self.verbose:
            logger.info(f"Semantic arXiv search: {query[:50]}...")
            if categories:
                logger.info(f"Category filter: {categories}")

        all_results = []

        search_queries = self._build_search_queries(query, categories)

        for sq in search_queries:
            try:
                results = self._search_with_query(
                    query=sq["query"],
                    max_results=max_results,
                    sort_by=sort_by,
                    field=sq.get("field", "all"),
                    from_year=from_year
                )
                all_results.extend(results)
            except Exception as e:
                if self.verbose:
                    logger.warning(f"Search failed for {sq['query']}: {e}")

        unique_results = self._deduplicate_results(all_results)
        scored_results = self._score_and_rank(unique_results, query)

        return scored_results[:max_results]

    def _build_search_queries(
        self,
        query: str,
        categories: Optional[list[str]] = None
    ) -> list[dict]:
        """Build multiple search queries for better coverage."""
        queries = []

        base_query = self._expand_query_from_text(query)

        if categories:
            cat_query = " OR ".join([f"cat:{c}" for c in categories])
            queries.append({
                "query": f"({base_query}) AND ({cat_query})",
                "field": "all"
            })
            queries.append({
                "query": f"({base_query}) AND ({cat_query})",
                "field": "title"
            })
        else:
            queries.append({
                "query": base_query,
                "field": "all"
            })

        return queries

    def _expand_query_from_text(self, query: str) -> str:
        """Extract keywords from query text and build semantic query."""
        extractor = SemanticKeywordExtractor(self.verbose)
        extracted = extractor.extract(query)

        terms = []
        keywords = extracted.get("keywords", [])
        bigrams = extracted.get("bigrams", [])[:5]

        for kw in keywords[:10]:
            if len(kw) > 3:
                terms.append(kw)

        terms.extend(bigrams)

        if terms:
            return f'({" OR ".join(terms)})'
        return query

    def _build_mbse_query(self) -> str:
        """Removed - now using dynamic keyword extraction."""
        return ""

    def _expand_query(self, query: str) -> str:
        """Removed - now using dynamic keyword extraction."""
        query_lower = query.lower()
        return query

    def _search_with_query(
        self,
        query: str,
        max_results: int,
        sort_by: str,
        field: str = "all",
        from_year: Optional[int] = None
    ) -> list[dict]:
        """Execute search with a specific query."""
        valid_sort_by = ["relevance", "lastUpdatedDate", "submittedDate"]
        valid_sort_order = ["ascending", "descending"]

        if sort_by not in valid_sort_by:
            sort_by = "relevance"

        search_field = f"{field}:{query}" if field != "all" else f"all:{query}"

        if from_year:
            search_field = f"({search_field}) AND submittedDate:[{from_year}-01-01 TO 9999-12-31]"

        params = {
            "search_query": search_field,
            "max_results": max_results,
            "sortBy": sort_by,
            "sortOrder": "descending"
        }

        try:
            url = f"{self.BASE_URL}?{urllib.parse.urlencode(params)}"
            response = self._make_request(url)

            entries = response.get("entry", [])
            normalized = self._normalize_entries(entries)

            for entry in normalized:
                entry["_search_query"] = query
                entry["_search_field"] = field

            return normalized

        except Exception as e:
            logger.error(f"arXiv search error: {e}")
            return []

    def _make_request(self, url: str, retries: int = 3) -> dict:
        """Make HTTP request to arXiv API."""
        last_error = None

        for attempt in range(retries):
            try:
                req = urllib.request.Request(url)
                req.add_header("User-Agent", "SemanticReview/1.0")

                with urllib.request.urlopen(req, timeout=30) as response:
                    import xml.etree.ElementTree as ET
                    data = response.read().decode("utf-8")

                    root = ET.fromstring(data)
                    ns = {"atom": "http://www.w3.org/2005/Atom", "arxiv": "http://arxiv.org/schemas/atom"}

                    entries = []
                    for entry in root.findall("atom:entry", ns):
                        entry_data = self._parse_entry(entry, ns)
                        entries.append(entry_data)

                    return {"entry": entries}

            except urllib.error.HTTPError as e:
                last_error = e
                if e.code == 500:
                    logger.warning(f"arXiv API 500 error (attempt {attempt + 1}/{retries})")
                    import time
                    time.sleep(self.request_delay * (attempt + 1))
                else:
                    raise ArxivApiError(f"HTTP error {e.code}: {e.reason}")
            except Exception as e:
                last_error = e
                logger.warning(f"Request error (attempt {attempt + 1}/{retries}): {e}")
                import time
                time.sleep(self.request_delay * (attempt + 1))

        raise ArxivApiError(f"Failed after {retries} attempts: {last_error}")

    def _parse_entry(self, entry, ns: dict) -> dict:
        """Parse XML entry to dictionary."""
        def find_text(element, path):
            el = element.find(path, ns)
            return el.text if el is not None else ""

        def find_category(element, term_key):
            categories = []
            for cat in element.findall("atom:category", ns):
                scheme = cat.get("scheme", "")
                if "arxiv.org" in scheme or term_key in scheme:
                    categories.append(cat.get("term", ""))
            return categories

        paper_data = {
            "id": find_text(entry, "atom:id").split("/abs")[-1],
            "title": find_text(entry, "atom:title").replace("\n", " ").strip(),
            "abstract": find_text(entry, "arxiv:summary").replace("\n", " ").strip(),
            "authors": [a.find("atom:name", ns).text for a in entry.findall("atom:author", ns)],
            "published": find_text(entry, "atom:published"),
            "updated": find_text(entry, "atom:updated"),
            "pdf_url": find_text(entry, "atom:link[@title='pdf']"),
            "categories": find_category(entry, "categories"),
            "comment": find_text(entry, "arxiv:comment"),
            "doi": find_text(entry, "arxiv:doi"),
            "journal_ref": find_text(entry, "arxiv:journal_ref"),
            "relevance_score": 0.0,
        }

        return paper_data

    def _normalize_entries(self, entries: list[dict]) -> list[dict]:
        """Normalize entries to consistent format."""
        return entries

    def _deduplicate_results(self, results: list[dict]) -> list[dict]:
        """Remove duplicate results based on arXiv ID."""
        seen = set()
        unique = []

        for result in results:
            paper_id = result.get("id", "")
            if paper_id and paper_id not in seen:
                seen.add(paper_id)
                unique.append(result)

        return unique

    EXCLUDED_PREFIXES = [
        "astro-ph", "gr-qc", "hep-", "nucl-", "cond-mat",
        "physics.", "math.", "q-bio", "nlin."
    ]

    EXCLUDED_EXACT = [
        "hep-ex", "hep-th", "hep-ph", "hep-lat",
        "nucl-th", "nucl-ex",
    ]

    def _is_cs_paper(self, categories: list[str]) -> bool:
        """Check if paper belongs to CS-related categories."""
        for cat in categories:
            for excluded in self.EXCLUDED_PREFIXES:
                if cat.startswith(excluded):
                    return False
            if cat in self.EXCLUDED_EXACT:
                return False
        return True

    def _score_and_rank(
        self,
        results: list[dict],
        original_query: str
    ) -> list[dict]:
        """Score and rank results by relevance to original query."""
        query_terms = self._extract_query_terms(original_query)

        scored_results = []
        excluded_count = 0

        for result in results:
            categories = result.get("categories", [])

            if not self._is_cs_paper(categories):
                excluded_count += 1
                continue

            score = self._calculate_relevance_score(result, query_terms)
            result["relevance_score"] = score
            scored_results.append(result)

        if self.verbose and excluded_count > 0:
            logger.info(f"Excluded {excluded_count} non-CS papers (physics/astronomy)")

        scored_results = sorted(
            scored_results,
            key=lambda x: x.get("relevance_score", 0),
            reverse=True
        )

        return scored_results

    def _extract_query_terms(self, query: str) -> list[str]:
        """Extract meaningful terms from query."""
        stop_words = {
            "a", "an", "the", "is", "are", "was", "were", "of", "in", "on",
            "for", "with", "and", "or", "to", "from", "by", "method", "approach"
        }
        terms = re.findall(r'\b[a-zA-Z]{3,}\b', query.lower())
        return [t for t in terms if t not in stop_words]

    def _calculate_relevance_score(
        self,
        result: dict,
        query_terms: list[str]
    ) -> float:
        """Calculate relevance score for a result."""
        if not query_terms:
            return 0.5

        title = result.get("title", "").lower()
        abstract = result.get("abstract", "").lower()
        categories = result.get("categories", [])

        title_score = sum(1 for term in query_terms if term in title) * 0.3
        abstract_score = sum(1 for term in query_terms if term in abstract) * 0.1

        category_boost = 0.0
        if any(cat in self.RELEVANT_CS_CATEGORIES for cat in categories):
            category_boost = 0.3

        recency_boost = 0.0
        try:
            year = int(result.get("published", "")[:4]) if result.get("published") else 0
            if year >= 2022:
                recency_boost = 0.1
        except:
            pass

        score = min(1.0, (title_score + abstract_score + category_boost + recency_boost))

        return round(score, 3)


class EnhancedOpenAlexSearcher:
    """Enhanced OpenAlex searcher with better filtering and scoring."""

    BASE_URL = "https://api.openalex.org/works"

    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        self.rate_limit_delay = 1.0

    def search(
        self,
        query: str,
        max_results: int = 20,
        domains: Optional[list[str]] = None,
        from_year: Optional[int] = None,
        to_year: Optional[int] = None
    ) -> list[dict]:
        """
        Search OpenAlex with enhanced filtering.
        
        Args:
            query: Search query
            max_results: Maximum results
            domains: Filter by domains (e.g., ["Computer Science", "Software Engineering"])
            from_year: Filter from year
            to_year: Filter to year
        """
        if self.verbose:
            logger.info(f"Enhanced OpenAlex search: {query[:50]}...")

        filters = []
        filters.append(f"title.search:{query}")
        filters.append(f"abstract.search:{query}")

        if from_year:
            filters.append(f"publication_year>={from_year}")
        if to_year:
            filters.append(f"publication_year<={to_year}")

        if domains:
            domain_filter = " OR ".join([f"concepts.display_name:{d}" for d in domains])
            filters.append(f"({domain_filter})")

        params = {
            "filter": ",".join(filters),
            "per_page": min(max_results, 50),
            "sort": "relevance_score:desc"
        }

        try:
            url = f"{self.BASE_URL}?{urllib.parse.urlencode(params)}"
            req = urllib.request.Request(url)
            req.add_header("User-Agent", "SemanticReview/1.0")

            import time
            time.sleep(self.rate_limit_delay)

            with urllib.request.urlopen(req, timeout=30) as response:
                data = json.loads(response.read().decode("utf-8"))
                results = [self._normalize_entry(item) for item in data.get("results", [])]

                scored_results = self._score_results(results, query)
                return scored_results

        except Exception as e:
            logger.error(f"OpenAlex search error: {e}")
            return []

    def _normalize_entry(self, entry: dict) -> dict:
        """Normalize OpenAlex entry to consistent format."""
        authors = []
        for author in entry.get("authorships", []):
            name = author.get("author", {}).get("display_name", "")
            if name:
                authors.append(name)

        return {
            "id": entry.get("doi", "").replace("https://doi.org/", ""),
            "title": entry.get("title", ""),
            "abstract": entry.get("abstract", "") or "",
            "authors": authors,
            "published": entry.get("publication_date", ""),
            "year": entry.get("publication_year"),
            "categories": [c.get("display_name", "") for c in entry.get("concepts", [])],
            "cited_by_count": entry.get("cited_by_count", 0),
            "url": entry.get("doi", ""),
            "relevance_score": 0.0,
        }

    def _score_results(
        self,
        results: list[dict],
        query: str
    ) -> list[dict]:
        """Score and rank results."""
        query_terms = self._extract_terms(query)

        for result in results:
            score = self._calculate_score(result, query_terms)
            result["relevance_score"] = score

        return sorted(results, key=lambda x: x.get("relevance_score", 0), reverse=True)

    def _extract_terms(self, query: str) -> list[str]:
        """Extract query terms."""
        stop_words = {"a", "an", "the", "method", "approach", "study"}
        terms = re.findall(r'\b[a-zA-Z]{3,}\b', query.lower())
        return [t for t in terms if t not in stop_words]

    def _calculate_score(self, result: dict, query_terms: list[str]) -> float:
        """Calculate relevance score."""
        if not query_terms:
            return 0.5

        title = result.get("title", "").lower()
        abstract = result.get("abstract", "").lower()

        title_hits = sum(1 for term in query_terms if term in title)
        abstract_hits = sum(1 for term in query_terms if term in abstract)

        score = min(1.0, title_hits * 0.25 + abstract_hits * 0.1)

        if result.get("cited_by_count", 0) > 100:
            score += 0.1

        return round(min(1.0, score), 3)


class ResearchDomainExtractor:
    """Extract research domain from paper for targeted search."""

    DOMAIN_KEYWORDS = {
        "software_engineering": [
            "software engineering", "software development", "code generation",
            "software architecture", "requirements", "testing", "refactoring",
            "software maintenance", "devops", "agile"
        ],
        "embedded_systems": [
            "embedded system", "real-time", "cyber-physical", "iot",
            "microcontroller", "firmware", "safety-critical", "rtos"
        ],
        "formal_methods": [
            "formal verification", "model checking", "theorem proving",
            "temporal logic", "formal specification", "safety verification"
        ],
        "programming_languages": [
            "programming language", "compiler", "type system", "semantics",
            "parser", "interpreter", "domain-specific language"
        ],
        "systems": [
            "distributed system", "operating system", "cloud computing",
            "network protocol", "distributed computing", "microservices"
        ],
        "ai_ml": [
            "machine learning", "deep learning", "neural network",
            "artificial intelligence", "natural language processing", "computer vision"
        ],
    }

    ARXIV_CATEGORIES = {
        "software_engineering": ["cs.SE", "cs.PL"],
        "embedded_systems": ["cs.OS", "cs.DC", "cs.SE"],
        "formal_methods": ["cs.LO", "cs.SE", "cs.FL"],
        "programming_languages": ["cs.PL", "cs.SE"],
        "systems": ["cs.OS", "cs.DC", "cs.NI"],
        "ai_ml": ["cs.AI", "cs.LG", "stat.ML"],
    }

    def extract(self, paper_data: dict) -> dict:
        """Extract research domain from paper."""
        text = f"{paper_data.get('title', '')} {paper_data.get('abstract', '')}".lower()

        domain_scores = {}
        for domain, keywords in self.DOMAIN_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in text)
            domain_scores[domain] = score

        primary_domain = max(domain_scores.items(), key=lambda x: x[1])
        relevant_domains = [
            d for d, s in domain_scores.items()
            if s > 0 and d != primary_domain[0]
        ] if primary_domain[1] > 0 else []

        return {
            "primary_domain": primary_domain[0] if primary_domain[1] > 0 else "general",
            "domain_score": primary_domain[1],
            "relevant_domains": relevant_domains,
            "arxiv_categories": self.ARXIV_CATEGORIES.get(primary_domain[0], ["cs.*"]),
        }


class EnhancedLiteratureSearcher:
    """Enhanced multi-source literature searcher with semantic search."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        self.arxiv_searcher = SemanticArxivSearcher(verbose)
        self.openalex_searcher = EnhancedOpenAlexSearcher(verbose)
        self.domain_extractor = ResearchDomainExtractor()
        self.keyword_extractor = SemanticKeywordExtractor(verbose)

    def search_papers(
        self,
        query: str,
        paper_data: Optional[dict] = None,
        core_count: int = 10,
        related_count: int = 15,
        background_count: int = 10
    ) -> dict[str, list[dict]]:
        """
        Search for literature with semantic understanding.
        
        Args:
            query: Base search query
            paper_data: Optional paper data to extract domain from
            core_count: Number of core (most relevant) papers
            related_count: Number of related papers
            background_count: Number of background papers
        """
        results = {
            "core": [],
            "related": [],
            "background": []
        }

        categories = None
        extracted_info = None

        if paper_data:
            domain_info = self.domain_extractor.extract(paper_data)
            categories = domain_info.get("arxiv_categories", ["cs.SE", "cs.OS"])

            paper_text = f"{paper_data.get('title', '')} {paper_data.get('abstract', '')}"
            extracted_info = self.keyword_extractor.extract(paper_text)

            if self.verbose:
                logger.info(f"Detected research domain: {domain_info['primary_domain']}")
                logger.info(f"Category filter: {categories}")
                logger.info(f"Extracted keywords: {extracted_info.get('keywords', [])[:10]}")

        if self.verbose:
            logger.info(f"Searching for core literature: {query[:50]}...")

        try:
            arxiv_results = self.arxiv_searcher.search(
                query=query,
                max_results=max(core_count * 3, 50),
                categories=categories,
                sort_by="relevance"
            )
            results["core"] = arxiv_results[:core_count]
        except Exception as e:
            logger.warning(f"Enhanced arXiv search failed: {e}")

        try:
            openalex_results = self.openalex_searcher.search(
                query=query,
                max_results=core_count // 2,
                from_year=2020
            )
            if openalex_results:
                existing_ids = {p.get("id", "") for p in results["core"]}
                for paper in openalex_results:
                    if paper.get("id") not in existing_ids:
                        results["core"].append(paper)
        except Exception as e:
            logger.warning(f"Enhanced OpenAlex search failed: {e}")

        results["core"] = results["core"][:core_count]

        if len(results["core"]) < 3 and extracted_info:
            if self.verbose:
                logger.info("Not enough CS papers found, generating dynamic fallback searches...")

            fallback_queries = self._generate_fallback_queries(extracted_info, query)

            for fq in fallback_queries[:5]:
                try:
                    fallback_results = self.arxiv_searcher.search(
                        query=fq,
                        max_results=10,
                        categories=categories or ["cs.SE", "cs.OS", "cs.PL"],
                        sort_by="relevance"
                    )
                    existing_ids = {p.get("id", "") for p in results["core"]}
                    for paper in fallback_results:
                        if paper.get("id") not in existing_ids and len(results["core"]) < core_count:
                            results["core"].append(paper)
                            existing_ids.add(paper.get("id", ""))
                except Exception as e:
                    if self.verbose:
                        logger.warning(f"Fallback search failed for '{fq[:30]}...': {e}")

        if self.verbose:
            logger.info(f"Found {len(results['core'])} core papers")

        if paper_data and categories:
            try:
                background_results = self.arxiv_searcher.search(
                    query=query,
                    max_results=background_count,
                    sort_by="submittedDate",
                    from_year=2018
                )
                existing_ids = {p.get("id", "") for p in results["core"]}
                background = [p for p in background_results if p.get("id") not in existing_ids]
                results["background"] = background[:background_count]
            except Exception as e:
                logger.warning(f"Background search failed: {e}")

        return results

    def _generate_fallback_queries(self, extracted_info: dict, original_query: str) -> list[str]:
        """Dynamically generate fallback queries from extracted keywords."""
        queries = []

        technical_terms = extracted_info.get("technical_terms", [])
        keywords = extracted_info.get("keywords", [])
        bigrams = extracted_info.get("bigrams", [])

        for term in technical_terms[:3]:
            if term.lower() not in original_query.lower():
                queries.append(term)

        for kw in keywords[:5]:
            if len(kw) > 3 and kw.lower() not in original_query.lower():
                combined = f'"{kw}" code generation'
                if combined not in queries:
                    queries.append(combined)

        for bg in bigrams[:3]:
            if bg.lower() not in original_query.lower():
                queries.append(bg)

        if not queries:
            domains = extracted_info.get("domains", [])
            for domain in domains:
                if domain == "software_engineering":
                    queries.append("software engineering code generation")
                elif domain == "embedded_systems":
                    queries.append("embedded system real-time software")
                elif domain == "ai_ml":
                    queries.append("machine learning code generation")
                elif domain == "mbse":
                    queries.append("model-based systems engineering")
                elif domain == "formal_methods":
                    queries.append("formal verification software")

        return queries[:5]


class LiteratureComparator:
    """Compare and analyze literature collections."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose

    def compare(
        self,
        target_paper: dict,
        literature: dict[str, list[dict]]
    ) -> dict:
        """Compare target paper against literature collection."""
        target_keywords = self._extract_keywords(target_paper)
        target_year = self._extract_year(target_paper)

        comparison = {
            "target_paper": {
                "title": target_paper.get("title", ""),
                "year": target_year,
                "keywords": target_keywords[:10],
            },
            "literature_summary": self._summarize_literature(literature),
            "novelty_analysis": self._analyze_novelty(target_paper, literature, target_keywords),
            "methodology_comparison": self._compare_methodology(target_paper, literature),
            "research_gaps": self._identify_research_gaps(target_paper, literature),
            "trends": self._identify_trends(literature),
        }

        return comparison

    def _extract_keywords(self, paper: dict) -> list[str]:
        """Extract keywords from paper."""
        text = f"{paper.get('title', '')} {paper.get('abstract', '')}"
        words = re.findall(r'\b[a-z]{4,}\b', text.lower())

        stop_words = {
            "this", "that", "with", "from", "have", "been", "were", "they",
            "their", "which", "more", "than", "also", "into", "some", "other"
        }

        keywords = [w for w in set(words) if w not in stop_words and len(w) > 4]

        return sorted(keywords, key=lambda x: words.count(x), reverse=True)[:20]

    def _extract_year(self, paper: dict) -> Optional[int]:
        """Extract publication year."""
        published = paper.get("published", "")
        if published:
            match = re.search(r'(\d{4})', published)
            if match:
                return int(match.group(1))
        return None

    def _summarize_literature(self, literature: dict) -> dict:
        """Summarize literature collection."""
        all_papers = literature.get("core", []) + literature.get("related", [])

        years = []
        for paper in all_papers:
            year = self._extract_year(paper)
            if year:
                years.append(year)

        return {
            "total_papers": len(all_papers),
            "year_range": {"min": min(years) if years else None, "max": max(years) if years else None},
            "avg_year": sum(years) / len(years) if years else None,
        }

    def _analyze_novelty(
        self,
        target_paper: dict,
        literature: dict,
        target_keywords: list[str]
    ) -> dict:
        """Analyze novelty of target paper."""
        all_papers = literature.get("core", []) + literature.get("related", [])

        literature_keywords = set()
        for paper in all_papers:
            keywords = self._extract_keywords(paper)
            literature_keywords.update(keywords[:10])

        unique_keywords = [kw for kw in target_keywords[:10] if kw not in literature_keywords]

        return {
            "unique_keywords": unique_keywords,
            "overlap_count": len(target_keywords) - len(unique_keywords),
            "novelty_assessment": "highly novel" if len(unique_keywords) > 5 else "moderately novel" if len(unique_keywords) > 2 else "incremental"
        }

    def _compare_methodology(self, target_paper: dict, literature: dict) -> dict:
        """Compare methodology with literature."""
        target_text = f"{target_paper.get('title', '')} {target_paper.get('abstract', '')}"

        target_methods = self._identify_methods(target_text)
        literature_methods = []

        for paper in literature.get("core", []):
            text = f"{paper.get('title', '')} {paper.get('abstract', '')}"
            literature_methods.extend(self._identify_methods(text))

        return {
            "target_methods": target_methods,
            "literature_methods": list(set(literature_methods))[:10],
            "comparison": "novel methods" if len(target_methods) > len(literature_methods) else "comparable methods"
        }

    def _identify_methods(self, text: str) -> list[str]:
        """Identify methods mentioned in text."""
        methods = []
        method_patterns = [
            r'\bneural network\b', r'\btransformer\b', r'\battention\b',
            r'\bgan\b', r'\bgpt\b', r'\bbert\b', r'\bresnet\b',
            r'\bgradient descent\b', r'\breinforcement learning\b',
            r'\bsupervised learning\b', r'\bunsupervised learning\b',
            r'\btransfer learning\b', r'\bfine-tuning\b'
        ]

        for pattern in method_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                methods.append(pattern.strip(r'\b'))

        return list(set(methods))

    def _identify_research_gaps(self, target_paper: dict, literature: dict) -> list[str]:
        """Identify research gaps."""
        gaps = []
        target_text = f"{target_paper.get('title', '')} {target_paper.get('abstract', '')}"

        gap_indicators = [
            (r'\blimitations?\b', "Limited by assumptions or constraints"),
            (r'\bfuture work\b', "Potential for extension"),
            (r'\bopen problem\b', "Open research problem"),
            (r'\bchallenge\b', "Ongoing challenges in the field"),
        ]

        for pattern, description in gap_indicators:
            if re.search(pattern, target_text, re.IGNORECASE):
                gaps.append(description)

        return gaps if gaps else ["General advancement in the field"]

    def _identify_trends(self, literature: dict) -> list[str]:
        """Identify research trends."""
        trends = []
        all_papers = literature.get("core", []) + literature.get("related", [])

        years = []
        for paper in all_papers:
            year = self._extract_year(paper)
            if year:
                years.append(year)

        if len(set(years)) > 1:
            sorted_years = sorted(years)
            if sorted_years[-1] - sorted_years[0] >= 3:
                trends.append("Long-term research trajectory")
                trends.append("Growing interest in recent years")

        trend_keywords = {
            "deep learning": ["neural", "deep", "network"],
            "transformers": ["transformer", "attention", "bert", "gpt"],
            "generative AI": ["generation", "gan", "diffusion"],
            "reinforcement": ["reinforcement", "policy", "agent"],
        }

        keyword_counts = {}
        for paper in all_papers:
            text = f"{paper.get('title', '')} {paper.get('abstract', '')}".lower()
            for trend, keywords in trend_keywords.items():
                count = sum(1 for kw in keywords if kw in text)
                if count > 0:
                    keyword_counts[trend] = keyword_counts.get(trend, 0) + count

        if keyword_counts:
            top_trend = max(keyword_counts.items(), key=lambda x: x[1])
            trends.append(f"Main trend: {top_trend[0]}")

        return trends


class ComprehensiveReviewWorkflow:
    """Main workflow orchestrator for comprehensive paper review."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        self.docx_parser = DocxParser(verbose)
        self.pdf_parser = PdfParser(verbose)
        self.literature_searcher = EnhancedLiteratureSearcher(verbose)
        self.literature_comparator = LiteratureComparator(verbose)

    def run(
        self,
        input_path: Optional[str] = None,
        arxiv_id: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> dict:
        """Execute the complete review workflow."""
        if self.verbose:
            logger.info("Starting Comprehensive Paper Review Workflow")

        result = {
            "workflow_version": "1.0",
            "timestamp": datetime.now().isoformat(),
            "phases": {}
        }

        try:
            if self.verbose:
                logger.info("Phase 1: Paper Input & Parsing")
            result["phases"]["input_parsing"] = self._phase1_input_parsing(input_path, arxiv_id)

            if self.verbose:
                logger.info("Phase 2: Literature Search & Collection")
            result["phases"]["literature_search"] = self._phase2_literature_search(result["phases"]["input_parsing"])

            if self.verbose:
                logger.info("Phase 3: Literature Analysis & Comparison")
            result["phases"]["literature_analysis"] = self._phase3_literature_analysis(
                result["phases"]["input_parsing"],
                result["phases"]["literature_search"]
            )

            if self.verbose:
                logger.info("Phase 4: Comprehensive Evaluation")
            result["phases"]["evaluation"] = self._phase4_evaluation(result["phases"]["input_parsing"])

            if self.verbose:
                logger.info("Phase 5: Peer Review")
            result["phases"]["peer_review"] = self._phase5_peer_review(
                result["phases"]["input_parsing"],
                result["phases"]["evaluation"]
            )

            if self.verbose:
                logger.info("Phase 6: Literature Comparison Report")
            result["phases"]["literature_report"] = self._phase6_literature_report(
                result["phases"]["input_parsing"],
                result["phases"]["literature_analysis"]
            )

            if self.verbose:
                logger.info("Phase 7: Final Review Report")
            result["phases"]["final_report"] = self._phase7_final_report(result)

            if output_path:
                self._save_report(result, output_path)

            return result

        except Exception as e:
            logger.error(f"Workflow error: {e}")
            result["error"] = str(e)
            return result

    def _phase1_input_parsing(
        self,
        input_path: Optional[str],
        arxiv_id: Optional[str]
    ) -> dict:
        """Phase 1: Parse input paper."""
        paper_data = None
        parse_error = None

        if arxiv_id:
            if self.verbose:
                logger.info(f"Fetching paper from arXiv: {arxiv_id}")
            try:
                paper_data = self.literature_searcher.arxiv_searcher.search_by_id(arxiv_id)
            except Exception as e:
                parse_error = str(e)

        if input_path:
            if self.verbose:
                logger.info(f"Parsing local document: {input_path}")
            try:
                input_file = Path(input_path)

                # Validate file exists
                if not input_file.exists():
                    raise FileNotFoundError(f"File not found: {input_path}")

                # Validate file is readable
                if not input_file.is_file():
                    raise ValueError(f"Path is not a file: {input_path}")

                ext = input_file.suffix.lower()
                if ext == ".docx":
                    parsed = self.docx_parser.parse(input_path)
                elif ext == ".pdf":
                    parsed = self.pdf_parser.parse(input_path)
                else:
                    raise ValueError(
                        f"Unsupported file format: {ext}. "
                        f"Supported formats: .docx, .pdf"
                    )

                paper_data = {
                    "title": parsed.metadata.title,
                    "abstract": parsed.metadata.abstract,
                    "source_type": parsed.metadata.source_type,
                    "arxiv_id": parsed.metadata.arxiv_id,
                    "parsed_sections": {
                        name: {
                            "title": section.title,
                            "content": section.content[:2000] if len(section.content) > 2000 else section.content
                        }
                        for name, section in parsed.sections.items()
                    },
                    "figures_count": len(parsed.figures),
                    "tables_count": len(parsed.tables),
                    "references_count": len(parsed.references),
                }

            except Exception as e:
                parse_error = str(e)

        if paper_data is None:
            paper_data = {}

        return {
            "paper_data": paper_data,
            "parse_error": parse_error,
            "success": parse_error is None
        }

    def _phase2_literature_search(self, phase1_result: dict) -> dict:
        """Phase 2: Search for related literature."""
        paper_data = phase1_result.get("paper_data", {})
        query = paper_data.get("title", "")

        if self.verbose:
            logger.info(f"Searching literature with query: {query[:50]}...")

        literature = {
            "core": [],
            "related": [],
            "background": []
        }

        if query:
            literature = self.literature_searcher.search_papers(query)

        return {
            "literature": literature,
            "core_count": len(literature.get("core", [])),
            "related_count": len(literature.get("related", [])),
            "background_count": len(literature.get("background", [])),
            "success": len(literature.get("core", [])) > 0 or len(literature.get("related", [])) > 0
        }

    def _phase3_literature_analysis(
        self,
        phase1_result: dict,
        phase2_result: dict
    ) -> dict:
        """Phase 3: Analyze and compare literature."""
        paper_data = phase1_result.get("paper_data", {})
        literature = phase2_result.get("literature", {})

        if self.verbose:
            logger.info("Analyzing literature collection")

        comparison = self.literature_comparator.compare(paper_data, literature)

        return {
            "comparison": comparison,
            "success": True
        }

    def _phase4_evaluation(self, phase1_result: dict) -> dict:
        """Phase 4: Comprehensive evaluation."""
        paper_data = phase1_result.get("paper_data", {})

        if self.verbose:
            logger.info("Evaluating paper")

        abstract = paper_data.get("abstract", "")
        title = paper_data.get("title", "")

        novelty = self._assess_novelty(title, abstract)
        methodology = self._assess_methodology(abstract)
        correctness = self._assess_correctness(paper_data)
        impact = self._assess_impact(paper_data)

        return {
            "innovation": novelty,
            "methodology": methodology,
            "correctness": correctness,
            "impact": impact,
            "overall_score": self._calculate_overall_score(novelty, methodology, correctness, impact),
            "success": True
        }

    def _phase5_peer_review(
        self,
        phase1_result: dict,
        phase4_result: dict
    ) -> dict:
        """Phase 5: Generate peer review."""
        paper_data = phase1_result.get("paper_data", {})
        evaluation = phase4_result

        if self.verbose:
            logger.info("Generating peer review")

        review = {
            "preliminary_assessment": self._preliminary_assessment(paper_data, evaluation),
            "section_review": self._section_review(paper_data, evaluation),
            "methodological_rigor": self._methodological_rigor(paper_data, evaluation),
            "reproducibility": self._reproducibility(paper_data),
            "ethics_compliance": self._ethics_compliance(),
            "writing_quality": self._writing_quality(paper_data),
            "final_recommendation": self._final_recommendation(evaluation)
        }

        return review

    def _phase6_literature_report(
        self,
        phase1_result: dict,
        phase3_result: dict
    ) -> dict:
        """Phase 6: Generate literature comparison report."""
        paper_data = phase1_result.get("paper_data", {})
        comparison = phase3_result.get("comparison", {})

        if self.verbose:
            logger.info("Generating literature comparison report")

        return {
            "current_state": comparison.get("literature_summary", {}),
            "novelty_claims": comparison.get("novelty_analysis", {}),
            "method_comparison": comparison.get("methodology_comparison", {}),
            "research_gaps": comparison.get("research_gaps", []),
            "trends": comparison.get("trends", []),
            "success": True
        }

    def _phase7_final_report(self, all_phases: dict) -> dict:
        """Phase 7: Generate final comprehensive report."""
        phase4 = all_phases.get("phases", {}).get("evaluation", {})
        phase5 = all_phases.get("phases", {}).get("peer_review", {})
        phase6 = all_phases.get("phases", {}).get("literature_report", {})

        overall_score = phase4.get("overall_score", 0)
        review = phase5.get("final_recommendation", {})

        if overall_score >= 0.7:
            recommendation = "Accept"
            confidence = "high"
        elif overall_score >= 0.5:
            recommendation = "Revise and Resubmit"
            confidence = "medium"
        else:
            recommendation = "Reject"
            confidence = "low"

        return {
            "executive_summary": self._generate_executive_summary(all_phases),
            "scores": {
                "innovation": phase4.get("innovation", {}).get("score", 0),
                "methodology": phase4.get("methodology", {}).get("score", 0),
                "correctness": phase4.get("correctness", {}).get("score", 0),
                "impact": phase4.get("impact", {}).get("score", 0),
                "overall": overall_score,
            },
            "recommendation": recommendation,
            "confidence": confidence,
            "revision_suggestions": phase5.get("preliminary_assessment", {}).get("suggestions", []),
            "strengths": self._identify_strengths(phase4),
            "weaknesses": self._identify_weaknesses(phase4),
            "success": True
        }

    def _assess_novelty(self, title: str, abstract: str) -> dict:
        """Assess paper novelty."""
        text = f"{title} {abstract}".lower()

        high_novelty = sum(1 for kw in [
            "first", "novel", "new", "introduce", "propose", "breakthrough"
        ] if kw in text)

        medium_novelty = sum(1 for kw in [
            "improve", "enhance", "extend", "combine", "integrate"
        ] if kw in text)

        score = min(1.0, (high_novelty * 0.25 + medium_novelty * 0.1))

        level = "high" if score >= 0.7 else "medium" if score >= 0.4 else "incremental"

        return {
            "score": round(score, 2),
            "level": level,
            "points": [kw for kw in ["first", "novel", "propose"] if kw in text][:3]
        }

    def _assess_methodology(self, abstract: str) -> dict:
        """Assess methodology quality."""
        text = abstract.lower()

        strengths = []
        for pattern in [r"experiment", r"evaluation", r"analysis", r"ablation", r"baseline"]:
            if re.search(pattern, text):
                strengths.append(pattern)

        score = min(1.0, 0.5 + len(strengths) * 0.1)

        return {
            "score": round(score, 2),
            "strengths": strengths[:5],
            "level": "strong" if score >= 0.7 else "moderate" if score >= 0.5 else "weak"
        }

    def _assess_correctness(self, paper_data: dict) -> dict:
        """Assess correctness claims."""
        return {
            "score": 0.7,
            "level": "likely correct",
            "notes": "Based on abstract analysis - full verification requires complete paper"
        }

    def _assess_impact(self, paper_data: dict) -> dict:
        """Assess potential impact."""
        categories = paper_data.get("categories", [])
        abstract = paper_data.get("abstract", "")

        impact_score = 0.5

        high_impact_cats = ["cs.AI", "cs.LG", "cs.CL", "stat.ML"]
        if any(cat in categories for cat in high_impact_cats):
            impact_score = 0.7

        return {
            "score": round(impact_score, 2),
            "level": "high" if impact_score >= 0.7 else "moderate",
            "domains": categories[:3]
        }

    def _calculate_overall_score(
        self,
        novelty: dict,
        methodology: dict,
        correctness: dict,
        impact: dict
    ) -> float:
        """Calculate overall score."""
        return round(
            novelty.get("score", 0.5) * 0.3 +
            methodology.get("score", 0.5) * 0.25 +
            correctness.get("score", 0.5) * 0.25 +
            impact.get("score", 0.5) * 0.2,
            2
        )

    def _preliminary_assessment(self, paper_data: dict, evaluation: dict) -> dict:
        """Generate preliminary assessment."""
        return {
            "summary": f"Paper addresses {paper_data.get('title', 'unknown')[:50]}...",
            "novelty_assessment": evaluation.get("innovation", {}).get("level", "unknown"),
            "impact_potential": "high" if evaluation.get("overall_score", 0) >= 0.7 else "moderate",
            "suitability": "Suitable for publication with revisions",
            "suggestions": [
                "Clarify methodology section",
                "Add more comparative experiments",
                "Strengthen novelty claims with evidence"
            ]
        }

    def _section_review(self, paper_data: dict, evaluation: dict) -> dict:
        """Generate section-by-section review."""
        return {
            "abstract": {"quality": "clear", "completeness": "complete"},
            "introduction": {"problem_clarity": "clear", "motivation": "adequate"},
            "methodology": {"technical_depth": "sufficient", "completeness": "complete"},
            "experiments": {"design": "rigorous", "baselines": "appropriate"},
            "results": {"presentation": "clear", "analysis": "thorough"},
            "discussion": {"interpretation": "reasonable", "limitations": "acknowledged"},
            "references": {"completeness": "complete", "relevance": "relevant"}
        }

    def _methodological_rigor(self, paper_data: dict, evaluation: dict) -> dict:
        """Assess methodological rigor."""
        return {
            "statistical_assumptions": "appropriate",
            "experimental_control": "adequate",
            "reproducibility": "high",
            "data_availability": "unclear",
            "code_availability": "unclear"
        }

    def _reproducibility(self, paper_data: dict) -> dict:
        """Assess reproducibility."""
        return {
            "method_description": "complete",
            "data_access": "unclear",
            "code_access": "unclear",
            "computational_requirements": "not specified"
        }

    def _ethics_compliance(self) -> dict:
        """Check ethics compliance."""
        return {
            "ethical_review": "Not applicable",
            "data_privacy": "Not specified",
            "consent": "Not specified"
        }

    def _writing_quality(self, paper_data: dict) -> dict:
        """Assess writing quality."""
        return {
            "clarity": "good",
            "organization": "logical",
            "grammar": "acceptable",
            "figures": "appropriate"
        }

    def _final_recommendation(self, evaluation: dict) -> dict:
        """Generate final recommendation."""
        score = evaluation.get("overall_score", 0)

        if score >= 0.7:
            recommendation = "Accept"
        elif score >= 0.5:
            recommendation = "Revise and Resubmit"
        else:
            recommendation = "Reject"

        return {
            "decision": recommendation,
            "confidence": "high" if score >= 0.7 else "medium",
            "summary": f"Overall score: {score}/1.0"
        }

    def _generate_executive_summary(self, all_phases: dict) -> str:
        """Generate executive summary."""
        phase4 = all_phases.get("phases", {}).get("evaluation", {})
        phase5 = all_phases.get("phases", {}).get("peer_review", {})
        phase1 = all_phases.get("phases", {}).get("input_parsing", {})

        paper_data = phase1.get("paper_data", {})
        title = paper_data.get("title", "Unknown paper")

        score = phase4.get("overall_score", 0)
        review = phase5.get("final_recommendation", {})

        return f"Review of '{title[:50]}...' - Overall Score: {score}/1.0, Recommendation: {review.get('decision', 'Pending')}"

    def _identify_strengths(self, evaluation: dict) -> list[str]:
        """Identify paper strengths."""
        strengths = []

        if evaluation.get("innovation", {}).get("score", 0) >= 0.7:
            strengths.append("Novel and impactful contribution")

        if evaluation.get("methodology", {}).get("score", 0) >= 0.7:
            strengths.append("Rigorous methodology")

        if evaluation.get("correctness", {}).get("score", 0) >= 0.7:
            strengths.append("Technically sound approach")

        return strengths if strengths else ["Well-structured presentation"]

    def _identify_weaknesses(self, evaluation: dict) -> list[str]:
        """Identify paper weaknesses."""
        weaknesses = []

        if evaluation.get("methodology", {}).get("score", 0) < 0.5:
            weaknesses.append("Methodology validation could be strengthened")

        if evaluation.get("innovation", {}).get("score", 0) < 0.5:
            weaknesses.append("Novelty claims need stronger evidence")

        if evaluation.get("impact", {}).get("score", 0) < 0.5:
            weaknesses.append("Impact could be more clearly articulated")

        return weaknesses if weaknesses else ["Minor suggestions for improvement"]

    def _save_report(self, result: dict, output_path: str):
        """Save report to file."""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        if self.verbose:
            logger.info(f"Report saved to: {output_path}")


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Comprehensive Paper Review Workflow - 7-phase peer review system"
    )
    parser.add_argument(
        "--input", "-i",
        help="Path to input document (DOCX or PDF)"
    )
    parser.add_argument(
        "--arxiv-id", "-a",
        help="arXiv paper ID (e.g., 2401.12345)"
    )
    parser.add_argument(
        "--output", "-o",
        help="Output file path for JSON report"
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable verbose logging"
    )

    args = parser.parse_args()

    if not args.input and not args.arxiv_id:
        parser.error("Either --input or --arxiv-id is required")

    workflow = ComprehensiveReviewWorkflow(verbose=args.verbose)

    result = workflow.run(
        input_path=args.input,
        arxiv_id=args.arxiv_id,
        output_path=args.output
    )

    if args.output:
        print(f"Report saved to: {args.output}")
    else:
        print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
