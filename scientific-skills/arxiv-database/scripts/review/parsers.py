"""
Paper parsers for different file formats.

Supports parsing DOCX, PDF documents and extracting structured
data including sections, figures, tables, and references.
"""

import logging
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional

from .data_models import (
    PaperFigure,
    PaperMetadata,
    PaperReference,
    PaperSection,
    PaperTable,
    StructuredPaperData,
)

# Check for optional dependencies
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pylatexenc.latex2text import LatexNodes2Text
    LATEX_AVAILABLE = True
except ImportError:
    LATEX_AVAILABLE = False


logger = logging.getLogger(__name__)


class PaperParser(ABC):
    """Abstract base class for paper parsers."""

    @abstractmethod
    def parse(self, source: str) -> StructuredPaperData:
        """Parse a paper from the given source."""
        pass


class DocxParser(PaperParser):
    """Parser for DOCX documents."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose

    def parse(self, file_path: str) -> StructuredPaperData:
        """Parse a DOCX document."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx>=1.0.0 "
                "or install all dependencies: pip install -e ."
            )

        doc = docx.Document(file_path)

        metadata = PaperMetadata(
            title=self._extract_title(doc),
            source_type="docx"
        )

        full_text = []
        sections = {}

        current_section = "Introduction"
        section_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text.append(text)

            section_header = self._detect_section_header(text)
            if section_header:
                if section_content:
                    sections[current_section] = PaperSection(
                        title=current_section,
                        content="\n".join(section_content)
                    )
                current_section = section_header
                section_content = []
            else:
                section_content.append(text)

        if section_content:
            sections[current_section] = PaperSection(
                title=current_section,
                content="\n".join(section_content)
            )

        abstract = sections.get("Abstract", PaperSection(title="Abstract", content=""))
        metadata.abstract = abstract.content if abstract else ""

        return StructuredPaperData(
            metadata=metadata,
            sections=sections,
            figures=self._extract_figures(doc),
            tables=self._extract_tables(doc),
            references=self._extract_references(doc)
        )

    def _extract_title(self, doc) -> str:
        """Extract title from document."""
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text.strip()
            if len(first_para) < 300:
                return first_para
        return "Unknown Title"

    def _detect_section_header(self, text: str) -> Optional[str]:
        """Detect if text is a section header."""
        patterns = [
            r'^(abstract)$',
            r'^1\.?\s*(introduction|overview)',
            r'^2\.?\s*(related work|background)',
            r'^3\.?\s*(method|approach|methodology)',
            r'^4\.?\s*(experiment|evaluation)',
            r'^5\.?\s*(result|results)',
            r'^6\.?\s*(discussion|conclusion)',
            r'^(acknowledgments?)',
            r'^(references|bibliography)',
        ]

        text_lower = text.lower().strip()
        for pattern in patterns:
            if re.match(pattern, text_lower):
                return text.title()

        if len(text) < 100 and text.isupper():
            return text.title()

        return None

    def _extract_figures(self, doc) -> list[PaperFigure]:
        """Extract figures from document."""
        figures = []
        pattern = r'(?:Figure|Fig\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n|$)'

        for para in doc.paragraphs:
            text = para.text.strip()
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                figures.append(PaperFigure(
                    number=match.group(1),
                    caption=match.group(2)[:200]
                ))

        return figures

    def _extract_tables(self, doc) -> list[PaperTable]:
        """Extract tables from document."""
        tables = []
        pattern = r'(?:Table|Tab\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n|$)'

        for para in doc.paragraphs:
            text = para.text.strip()
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                tables.append(PaperTable(
                    number=match.group(1),
                    caption=match.group(2)[:200]
                ))

        return tables

    def _extract_references(self, doc) -> list[PaperReference]:
        """Extract references from document."""
        references = []
        ref_pattern = r'\[(\d+)\]\s*(.+?)(?=\n\s*\[|\n\s*\d+\.|$)'
        in_references = False

        for para in doc.paragraphs:
            text = para.text.strip()

            if re.match(r'^(references|bibliography)$', text.lower()):
                in_references = True
                continue

            if in_references:
                match = re.match(ref_pattern, text, re.DOTALL)
                if match:
                    ref_text = match.group(2).strip()
                    arxiv_id = self._extract_arxiv_id(ref_text)
                    references.append(PaperReference(
                        id=match.group(1),
                        text=ref_text[:500],
                        arxiv_id=arxiv_id
                    ))

        return references

    def _extract_arxiv_id(self, text: str) -> Optional[str]:
        """Extract arXiv ID from text."""
        match = re.search(r'arXiv[:\.]?\s*(\d+\.\d+)', text, re.IGNORECASE)
        return match.group(1) if match else None


class PdfParser(PaperParser):
    """Parser for PDF documents using pypdf."""

    def __init__(self, verbose: bool = False):
        self.verbose = verbose

    def parse(self, file_path: str) -> StructuredPaperData:
        """Parse a PDF document."""
        if not PDF_AVAILABLE:
            raise ImportError(
                "pypdf is required for PDF parsing. "
                "Install with: pip install pypdf>=3.0.0 "
                "or install all dependencies: pip install -e ."
            )

        reader = pypdf.PdfReader(file_path)
        full_text = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)

        text = "\n\n".join(full_text)

        sections = self._extract_sections(text)
        metadata = self._extract_metadata(text, file_path)
        metadata.source_type = "pdf"

        abstract = sections.get("Abstract", PaperSection(title="Abstract", content=""))
        metadata.abstract = abstract.content if abstract else ""

        return StructuredPaperData(
            metadata=metadata,
            sections=sections,
            figures=self._extract_figures(text),
            tables=self._extract_tables(text),
            references=self._extract_references(text)
        )

    def _extract_sections(self, text: str) -> dict[str, PaperSection]:
        """Extract sections from text."""
        sections = {}

        section_patterns = {
            "abstract": r'(?i)(?:^|\n)(abstract)\s*[:\n](.+?)(?=\n\s*(?:1\.?|introduction|background|related work))',
            "introduction": r'(?i)(?:^|\n)(1\s*\.?\s*(?:introduction)?)\s*[:\n](.+?)(?=\n\s*(?:2\.?|related work|background))',
            "related_work": r'(?i)(?:^|\n)(2\s*\.?\s*(?:related work|background)?)\s*[:\n](.+?)(?=\n\s*(?:3\.?|method|approach|methodology))',
            "methodology": r'(?i)(?:^|\n)(3\s*\.?\s*(?:method|approach|methodology)?)\s*[:\n](.+?)(?=\n\s*(?:4\.?|experiment|evaluation|results|discussion))',
            "experiments": r'(?i)(?:^|\n)(4\s*\.?\s*(?:experiment|evaluation|experiments)?)\s*[:\n](.+?)(?=\n\s*(?:5\.?|result|discussion|conclusion))',
            "results": r'(?i)(?:^|\n)(5\s*\.?\s*(?:result|results|discussion)?)\s*[:\n](.+?)(?=\n\s*(?:6\.?|conclusion|future|limitation))',
            "conclusion": r'(?i)(?:^|\n)(6\s*\.?\s*(?:conclusion|conclusions|discussion)?)\s*[:\n](.+?)(?=\n\s*(?:references|bibliography|$))',
            "acknowledgments": r'(?i)(?:^|\n)(acknowledgments?)\s*[:\n](.+?)(?=\n\s*(?:references|bibliography|$))',
        }

        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.DOTALL)
            if match:
                content = match.group(2).strip()
                if len(content) > 10000:
                    content = content[:10000] + "..."
                sections[section_name] = PaperSection(
                    title=section_name.title(),
                    content=content
                )

        return sections

    def _extract_metadata(self, text: str, file_path: str) -> PaperMetadata:
        """Extract metadata from PDF text."""
        return PaperMetadata(
            title=self._extract_title(text, file_path),
            abstract="",
            source_type="pdf"
        )

    def _extract_title(self, text: str, file_path: str) -> str:
        """Extract title from text."""
        lines = text.split("\n")
        for line in lines[:10]:
            line = line.strip()
            if line and len(line) > 10 and len(line) < 300:
                if not line.startswith("http"):
                    return line
        return Path(file_path).stem

    def _extract_figures(self, text: str) -> list[PaperFigure]:
        """Extract figures from text."""
        figures = []
        pattern = r'(?:Figure|Fig\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n\s*(?:Figure|Fig\.?|Table|$))'

        for match in re.finditer(pattern, text, re.IGNORECASE):
            figures.append(PaperFigure(
                number=match.group(1),
                caption=match.group(2).strip()[:200]
            ))

        return figures

    def _extract_tables(self, text: str) -> list[PaperTable]:
        """Extract tables from text."""
        tables = []
        pattern = r'(?:Table|Tab\.?)\s*(\d+)[:\.\s]+(.+?)(?=\n\s*(?:Table|Tab\.?|Figure|$))'

        for match in re.finditer(pattern, text, re.IGNORECASE):
            tables.append(PaperTable(
                number=match.group(1),
                caption=match.group(2).strip()[:200]
            ))

        return tables

    def _extract_references(self, text: str) -> list[PaperReference]:
        """Extract references from text."""
        references = []
        ref_match = re.search(r'(?i)(?:references|bibliography)\s*[:\n](.+)', text, re.DOTALL)
        if not ref_match:
            return references

        ref_text = ref_match.group(1)
        bracket_pattern = r'\[(\d+)\]\s*(.+?)(?=\n\s*\[|\n\s*\d+\.|\Z)'

        for match in re.finditer(bracket_pattern, ref_text, re.DOTALL):
            ref_content = match.group(2).strip()
            arxiv_id = self._extract_arxiv_id(ref_content)
            references.append(PaperReference(
                id=match.group(1),
                text=ref_content[:500],
                arxiv_id=arxiv_id
            ))

        return references

    def _extract_arxiv_id(self, text: str) -> Optional[str]:
        """Extract arXiv ID from text."""
        match = re.search(r'arXiv[:\.]?\s*(\d+\.\d+)', text, re.IGNORECASE)
        return match.group(1) if match else None


class LatexParser(PaperParser):
    """
    Parser for LaTeX documents (.tex files).

    Extracts structured information while preserving formatting context.
    Requires pylatexenc package for LaTeX-to-text conversion.
    """

    def __init__(self, verbose: bool = False):
        """
        Initialize LaTeX parser.

        Args:
            verbose: Enable verbose logging
        """
        if not LATEX_AVAILABLE:
            raise ImportError(
                "pylatexenc not available. LaTeX parsing requires pylatexenc. "
                "Install with: pip install pylatexenc>=2.10"
            )

        self.verbose = verbose
        self.latex_converter = LatexNodes2Text()

        if verbose:
            logger.setLevel(logging.DEBUG)

    def parse(self, file_path: str) -> StructuredPaperData:
        """
        Parse LaTeX file and extract structured data.

        Args:
            file_path: Path to .tex file

        Returns:
            StructuredPaperData containing parsed information

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a LaTeX file
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if path.suffix.lower() not in ['.tex', '.latex']:
            raise ValueError(f"Not a LaTeX file: {file_path}")

        logger.info(f"Parsing LaTeX file: {path.name}")

        # Read LaTeX source
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            latex_content = f.read()

        # Extract metadata
        metadata = self._extract_metadata(latex_content)

        # Extract sections
        sections = self._extract_sections(latex_content)

        # Extract figures
        figures = self._extract_figures(latex_content)

        # Extract tables
        tables = self._extract_tables(latex_content)

        # Extract references
        references = self._extract_references(latex_content)

        return StructuredPaperData(
            metadata=metadata,
            sections=sections,
            figures=figures,
            tables=tables,
            references=references
        )

    def _extract_metadata(self, latex_content: str) -> PaperMetadata:
        """Extract metadata from LaTeX preamble."""
        metadata = PaperMetadata(title="", source_type="latex")

        # Extract title
        title_match = re.search(
            r'\\title\s*\{([^}]+)\}',
            latex_content,
            re.IGNORECASE | re.DOTALL
        )
        if title_match:
            title_latex = title_match.group(1)
            metadata.title = self.latex_converter.latex_to_text(title_latex).strip()

        # Extract authors
        author_match = re.search(
            r'\\author\s*\{([^}]+)\}',
            latex_content,
            re.IGNORECASE | re.DOTALL
        )
        if author_match:
            author_latex = author_match.group(1)
            author_text = self.latex_converter.latex_to_text(author_latex)
            # Split by common separators
            authors = re.split(r'\s+and\s+|,\s*|\n', author_text)
            metadata.authors = [a.strip() for a in authors if a.strip()]

        # Extract abstract
        abstract_match = re.search(
            r'\\begin\{abstract\}(.*?)\\end\{abstract\}',
            latex_content,
            re.IGNORECASE | re.DOTALL
        )
        if abstract_match:
            abstract_latex = abstract_match.group(1)
            metadata.abstract = self.latex_converter.latex_to_text(abstract_latex).strip()

        return metadata

    def _extract_sections(self, latex_content: str) -> dict[str, PaperSection]:
        """Extract sections from LaTeX content."""
        sections = {}

        # Section patterns (in order of hierarchy)
        section_patterns = [
            (r'\\section\s*\{([^}]+)\}(.*?)(?=\\section|\\subsection|\\bibliography|\\end\{document\}|\Z)',
             1),
            (r'\\subsection\s*\{([^}]+)\}(.*?)(?=\\section|\\subsection|\\subsubsection|\\bibliography|\\end\{document\}|\Z)',
             2),
            (r'\\subsubsection\s*\{([^}]+)\}(.*?)(?=\\section|\\subsection|\\subsubsection|\\bibliography|\\end\{document\}|\Z)',
             3),
        ]

        for pattern, level in section_patterns:
            for match in re.finditer(pattern, latex_content, re.DOTALL | re.IGNORECASE):
                section_title_latex = match.group(1)
                section_content_latex = match.group(2)

                # Convert LaTeX to text
                section_title = self.latex_converter.latex_to_text(section_title_latex).strip()
                section_content = self.latex_converter.latex_to_text(section_content_latex).strip()

                # Clean up content
                section_content = re.sub(r'\n\s*\n', '\n\n', section_content)

                sections[section_title] = PaperSection(
                    title=section_title,
                    content=section_content,
                    level=level
                )

        return sections

    def _extract_figures(self, latex_content: str) -> list[PaperFigure]:
        """Extract figures from LaTeX content."""
        figures = []

        # Find all figure environments
        figure_pattern = r'\\begin\{figure\}(.*?)\\end\{figure\}'

        for match in re.finditer(figure_pattern, latex_content, re.DOTALL | re.IGNORECASE):
            figure_content = match.group(1)

            # Extract caption
            caption_match = re.search(
                r'\\caption\s*\{([^}]+)\}',
                figure_content,
                re.IGNORECASE | re.DOTALL
            )

            if caption_match:
                caption_latex = caption_match.group(1)
                caption = self.latex_converter.latex_to_text(caption_latex).strip()

                # Try to extract figure number from label or guess
                label_match = re.search(r'\\label\s*\{fig:([^}]+)\}', figure_content)
                fig_number = label_match.group(1) if label_match else str(len(figures) + 1)

                figures.append(PaperFigure(
                    number=fig_number,
                    caption=caption
                ))

        return figures

    def _extract_tables(self, latex_content: str) -> list[PaperTable]:
        """Extract tables from LaTeX content."""
        tables = []

        # Find all table environments
        table_pattern = r'\\begin\{table\}(.*?)\\end\{table\}'

        for match in re.finditer(table_pattern, latex_content, re.DOTALL | re.IGNORECASE):
            table_content = match.group(1)

            # Extract caption
            caption_match = re.search(
                r'\\caption\s*\{([^}]+)\}',
                table_content,
                re.IGNORECASE | re.DOTALL
            )

            if caption_match:
                caption_latex = caption_match.group(1)
                caption = self.latex_converter.latex_to_text(caption_latex).strip()

                # Try to extract table number from label or guess
                label_match = re.search(r'\\label\s*\{tab:([^}]+)\}', table_content)
                table_number = label_match.group(1) if label_match else str(len(tables) + 1)

                tables.append(PaperTable(
                    number=table_number,
                    caption=caption
                ))

        return tables

    def _extract_references(self, latex_content: str) -> list[PaperReference]:
        """Extract references from LaTeX bibliography."""
        references = []

        # Find bibliography section
        bib_pattern = r'\\begin\{thebibliography\}(.*?)\\end\{thebibliography\}'
        bib_match = re.search(bib_pattern, latex_content, re.DOTALL | re.IGNORECASE)

        if bib_match:
            bib_content = bib_match.group(1)

            # Extract individual bibitem entries
            bibitem_pattern = r'\\bibitem\s*\{([^}]+)\}\s*([^\\]+)'

            for match in re.finditer(bibitem_pattern, bib_content, re.DOTALL):
                ref_id = match.group(1)
                ref_text_latex = match.group(2)

                # Convert to text
                ref_text = self.latex_converter.latex_to_text(ref_text_latex).strip()

                # Extract arXiv ID if present
                arxiv_id = self._extract_arxiv_id(ref_text)

                # Extract DOI if present
                doi = self._extract_doi(ref_text)

                references.append(PaperReference(
                    id=ref_id,
                    text=ref_text[:500],  # Limit length
                    arxiv_id=arxiv_id,
                    doi=doi
                ))

        # Also check for external .bib file references
        bibliography_match = re.search(r'\\bibliography\s*\{([^}]+)\}', latex_content)
        if bibliography_match and not references:
            bib_file = bibliography_match.group(1)
            logger.warning(
                f"Found external bibliography file: {bib_file}.bib. "
                "External .bib files are not parsed automatically. "
                "Consider using BibTeX tools separately."
            )

        return references

    def _extract_arxiv_id(self, text: str) -> Optional[str]:
        """Extract arXiv ID from text."""
        match = re.search(r'arXiv[:\.]?\s*(\d+\.\d+)', text, re.IGNORECASE)
        return match.group(1) if match else None

    def _extract_doi(self, text: str) -> Optional[str]:
        """Extract DOI from text."""
        match = re.search(r'doi[:\.]?\s*(10\.\d+/[^\s]+)', text, re.IGNORECASE)
        return match.group(1) if match else None
